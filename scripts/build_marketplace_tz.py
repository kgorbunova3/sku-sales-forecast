"""Генерує ТЗ_marketplace.docx за шаблоном викладача."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "TZ_marketplace.docx"


def add_heading(doc, text, level=1, size=14, bold=True, color=(0, 0, 0)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(8 if level > 1 else 14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def build():
    doc = Document()

    # Стиль документа
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Титул
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ТЕХНІЧНЕ ЗАВДАННЯ")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("на розробку алгоритму прогнозування продажів у розрізі SKU")
    run.italic = True
    run.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Виконавець: Горбунова Крістіна, група УМ-з31\n").italic = True
    meta.add_run("Травень 2026").italic = True

    doc.add_paragraph()  # spacer

    # ------------------------------------------------------------------ 1
    add_heading(doc, "1. Назва проекту / задачі")
    add_body(
        doc,
        "Прогнозування тижневих продажів товарів у розрізі SKU × магазин на горизонт "
        "4 тижні наперед для платформи міжнародного маркетплейсу.",
    )

    # ------------------------------------------------------------------ 2
    add_heading(doc, "2. Тип моделі")
    add_bullets(
        doc,
        [
            "Основна: Gradient Boosting (LightGBM) на flat-таблиці з календарними та лаговими ознаками.",
            "Baseline: наївний прогноз = середнє за 4 останні тижні (для порівняння та fallback).",
            "Стратегія прогнозу: рекурсивний multi-step (на горизонт 4 тижні).",
        ],
    )

    # ------------------------------------------------------------------ 3
    add_heading(doc, "3. Посилання на таски (Kanban / Jira)")
    add_body(
        doc,
        "Робота поділена на 5 задач у GitHub Project (Kanban-дошка проекту):",
    )
    add_kv(doc, "Kanban-дошка", "https://github.com/users/kgorbunova3/projects/1")
    add_bullets(
        doc,
        [
            "Задача #1 — Написання ТЗ: https://github.com/kgorbunova3/sku-sales-forecast/issues/1",
            "Задача #2 — Формування Датасету: https://github.com/kgorbunova3/sku-sales-forecast/issues/2",
            "Задача #3 — Вибір та запуск моделі: https://github.com/kgorbunova3/sku-sales-forecast/issues/3",
            "Задача #4 — Вибір та оцінка метрик: https://github.com/kgorbunova3/sku-sales-forecast/issues/4",
            "Задача #5 — Формування звіту з результатами: https://github.com/kgorbunova3/sku-sales-forecast/issues/5",
        ],
    )

    # ------------------------------------------------------------------ 4
    add_heading(doc, "4. Дата запуску проекту в продакшн")
    add_body(doc, "Планова дата інтеграції моделі у робочий процес: 30 червня 2026 (кінець семестру).")

    # ------------------------------------------------------------------ 5
    add_heading(doc, "5. Розміщення проекту")
    add_kv(doc, "Репозиторій", "https://github.com/kgorbunova3/sku-sales-forecast")
    add_kv(
        doc,
        "Джерела вхідних даних",
        "внутрішні таблиці маркетплейсу (DWH BigQuery / Snowflake):",
    )
    add_bullets(
        doc,
        [
            "fact_sales_daily — щоденні продажі (date, store_id, sku_id, qty, gmv)",
            "dim_sku — каталог SKU (sku_id, category, brand, price_avg)",
            "dim_store — довідник магазинів / fulfillment-центрів",
            "fact_promo_calendar — календар акцій і знижок",
            "fact_inventory_snapshot — щоденні залишки та товар у дорозі (GiT)",
        ],
    )
    add_kv(
        doc,
        "Результати",
        "вихід моделі зберігається у таблиці forecast_sku_weekly у BigQuery (партиція по run_date) "
        "та паралельно у S3 у форматі Parquet: s3://demand-forecast/output/forecast_sku/run_date=YYYY-MM-DD/",
    )
    add_body(
        doc,
        "Формат запису: week_start_date, store_id, sku_id, yhat, yhat_lower, yhat_upper, model_version, run_ts.",
    )
    add_kv(
        doc,
        "Передача даних у інші системи",
        "прогноз транслюється у систему планування закупок (Replenishment) "
        "та у BI-дашборд Tableau «Demand Planning».",
    )

    # ------------------------------------------------------------------ 6
    add_heading(doc, "6. Регулярність запуску")
    add_bullets(
        doc,
        [
            "Щотижня, понеділок 06:00 UTC (Airflow DAG demand_forecast_sku_weekly).",
            "Перенавчання моделі — раз на квартал або при WAPE > 30% дві неділі поспіль.",
            "Виправний ре-запуск (re-run) при оновленні даних — у напівавтоматичному режимі.",
        ],
    )

    # ------------------------------------------------------------------ 7
    add_heading(doc, "7. Замовник")
    add_body(
        doc,
        "Department of Supply Chain & Demand Planning (відділ управління запасами).\n"
        "Бізнес-власник: Head of Replenishment.\n"
        "Внутрішній споживач прогнозу: Inventory Planners (~25 ос.), Category Managers (~40 ос.).",
    )

    # ------------------------------------------------------------------ 8
    add_heading(doc, "8. Розробник(и)")
    add_bullets(
        doc,
        [
            "Data Scientist (виконавець): Горбунова Крістіна (УМ-з31).",
            "Data Engineer (інтеграції BigQuery/Airflow): Demand Forecast Squad.",
            "Data Analyst (валідація бізнес-метрик): Replenishment Team.",
        ],
    )

    # ------------------------------------------------------------------ 9
    add_heading(doc, "9. Яка проблема вирішується")
    add_body(
        doc,
        "Маркетплейс зазнає одночасних втрат із двох сторін:",
    )
    add_bullets(
        doc,
        [
            "Out-of-Stock (OOS): потенційний клієнт не знаходить товар і йде до конкурента — пряма "
            "втрата виручки + ризик зниження customer retention.",
            "Over-Stock: на складах накопичується товар, який не продається — заморожуються оборотні "
            "кошти, ростуть витрати на зберігання, ризик уцінки.",
            "Ручне планування Inventory Planner'ами не масштабується на десятки тисяч SKU.",
        ],
    )
    add_body(
        doc,
        "Алгоритм має автоматично генерувати тижневі прогнози продажів у розрізі SKU × магазин, "
        "які лягають в основу формули обсягу замовлення:",
    )
    add_kv(doc, "Формула замовлення", "Q = Fct − St − GiT + SS_Fct")
    add_body(
        doc,
        "де Fct — прогноз продажів на горизонт, St — поточний залишок, GiT — товар у дорозі, "
        "SS_Fct — страховий запас (≈ 1.65·σ для service level 95%).",
        italic=True,
    )

    # ------------------------------------------------------------------ 10
    add_heading(doc, "10. Бізнес-вимоги до проекту")
    add_bullets(
        doc,
        [
            "Прогноз тижневих продажів на 4 тижні вперед у розрізі SKU × магазин (повна матриця).",
            "Прогноз має поставлятись у BI-систему Tableau та CSV/Parquet-експорт.",
            "Latency розрахунку: повний запуск (15 SKU × 2 магазини × 4 тижні) — менш ніж 5 хв.",
            "Стабільність: повторні запуски на тих самих даних дають детерміновані результати.",
            "Можливість фільтру за категорією / магазином / SKU у дашборді.",
        ],
    )

    # ------------------------------------------------------------------ 11
    add_heading(doc, "11. Метрики якості")
    add_kv(doc, "Основна метрика", "WAPE (Weighted Absolute Percentage Error)")
    add_kv(doc, "Поріг прийнятності", "WAPE ≤ 20% на тестовому періоді (останні 8 тижнів)")
    add_bullets(
        doc,
        [
            "MAPE — допоміжна метрика (контроль викидів)",
            "MAE / RMSE — для порівняння категорій",
            "Bias — у межах ±5% (систематична помилка прогнозу)",
            "SLA latency — менше 5 хв на повний запуск",
        ],
    )
    add_body(
        doc,
        "Поточні результати моделі (LightGBM): WAPE 25.24%, MAE 9.61 — тобто модель перевищує "
        "поточний поріг і потребує подальшого тюнінгу (вище порогу через малий синтетичний датасет; "
        "у продакшні з повним обсягом даних очікується WAPE ≤ 20%).",
        italic=True,
    )

    # ------------------------------------------------------------------ 12
    add_heading(doc, "12. Виключення з прогнозу / обмеження")
    add_bullets(
        doc,
        [
            "Не прогнозуються SKU зі статусом out-of-stock понад 8 тижнів (cold start).",
            "Не прогнозуються товари, знятi з асортименту (lifecycle = 'discontinued').",
            "Сезонні промо (Black Friday, Cyber Monday) обробляються окремою фічею is_holiday_week, "
            "але рідкісні «чорні лебеді» (геополітика, локдаун) можуть давати помилку >50%.",
            "Новi SKU (історія < 12 тижнів) — використовується fallback на категорійне середнє.",
        ],
    )

    # ------------------------------------------------------------------ 13
    add_heading(doc, "13. Особливості / обмеження")
    add_bullets(
        doc,
        [
            "Технічні: модель навчається на pandas + LightGBM, ресурс — 1 CPU-узел, RAM ≤ 8 ГБ.",
            "Бізнес: горизонт 4 тижні фіксований (узгоджено з циклом замовлення постачальникам).",
            "Дані оновлюються щодня, але модель агрегує до тижня (тижні з понеділка, W-MON).",
            "Прогноз — мінімум 0 (без негативних значень), без обмеження зверху.",
        ],
    )

    # ------------------------------------------------------------------ 14
    add_heading(doc, "14. Економічна ефективність")
    add_bullets(
        doc,
        [
            "Зниження OOS-rate з очікуваних 12–15% до 7–8% — потенційне зростання GMV на ~3% по "
            "пілотних категоріях.",
            "Скорочення over-stock на 10–15% — звільнення оборотних коштів орієнтовно ~150 тис. USD "
            "по пілоту з 15 SKU.",
            "Економія часу Inventory Planner'ів: автоматизація рутинного планування ~20 годин/тиждень "
            "на 1 категорію.",
            "Відмова від платних сторонніх SaaS-сервісів прогнозування (типу o9, RELEX) — "
            "економія ліцензій від $50k/рік.",
        ],
    )

    # ------------------------------------------------------------------ 15
    add_heading(doc, "15. Презентація")
    add_body(
        doc,
        "У розробці. Посилання на слайди для захисту проекту буде додано до фінального дедлайну "
        "(30 червня 2026). Місце збереження: спільний Google Drive команди.",
        italic=True,
    )

    # ------------------------------------------------------------------ 16
    add_heading(doc, "16. Детальний опис проекту")

    add_heading(doc, "16.1. Функціональні вимоги", level=2, size=12)
    add_bullets(
        doc,
        [
            "Зчитувати свіжі денні продажі з fact_sales_daily за останні 12 місяців.",
            "Агрегувати дані до тижневої гранулярності (W-MON).",
            "Будувати ознаки: лаги (1, 2, 4, 12), ковзні середні, календарні (woy_sin/cos, holiday).",
            "Тренувати LightGBM (з early stopping) і генерувати прогноз на 4 тижні наперед.",
            "Виводити прогноз у таблицю forecast_sku_weekly та у CSV-експорт для дашборду.",
            "Логувати метрики моделі та feature importance у MLflow / W&B.",
            "Викликати API replenishment-system для трансляції плану замовлення Q.",
        ],
    )

    add_heading(doc, "16.2. Flow роботи (покрокова схема)", level=2, size=12)
    add_bullets(
        doc,
        [
            "1) Airflow DAG demand_forecast_sku_weekly стартує по cron щопонеділка 06:00 UTC.",
            "2) Task fetch_data — SQL-запит до DWH, вивантаження продажів, залишків, промо.",
            "3) Task aggregate_weekly — агрегація денних даних до тижневих.",
            "4) Task feature_engineering — генерація лагів, ковзних, календарних ознак.",
            "5) Task train_or_load_model — інкрементне навчання або завантаження поточної моделі з MLflow.",
            "6) Task predict — рекурсивний прогноз на 4 тижні.",
            "7) Task compute_order_qty — застосування формули Q = Fct − St − GiT + SS_Fct.",
            "8) Task export_results — запис у BigQuery + Parquet у S3.",
            "9) Task notify — Slack-нотифікація команди Replenishment про готовність.",
        ],
    )

    add_heading(doc, "16.3. Блок-схема", level=2, size=12)
    add_body(
        doc,
        "[DWH BigQuery] → fetch_data → aggregate_weekly → feature_engineering → "
        "train/load_model → predict → compute_order_qty → "
        "[BigQuery forecast_sku_weekly] + [S3 Parquet] → Tableau dashboard + Replenishment API → Slack notify",
        italic=True,
    )

    # ------------------------------------------------------------------ 17
    add_heading(doc, "17. Вимоги до даних")
    add_body(doc, "Структура вхідного датасету (один рядок = день × магазин × SKU):")
    add_bullets(
        doc,
        [
            "date — date (ISO)",
            "store_id — int",
            "sku_id — int",
            "category — string",
            "qty — int (кількість одиниць)",
            "price — float (грн)",
            "promo — int (0/1)",
        ],
    )
    add_bullets(
        doc,
        [
            "Оновлення: щодня о 02:00 UTC (попередня доба + correction-batch за 7 днів назад).",
            "Формат у DWH: BigQuery partitioned by date; у S3 — Parquet snappy.",
            "Інтеграція: через Airflow operators (BigQueryHook, S3Hook).",
            "Якість: data quality checks — null rate qty < 1%, відсутність duplicates по (date, store, sku).",
        ],
    )

    # ------------------------------------------------------------------ 18
    add_heading(doc, "18. Інтерфейс користувача")
    add_body(
        doc,
        "Власний UI не передбачений. Користувачі (Inventory Planners, Category Managers) "
        "взаємодіють з результатами через:",
    )
    add_bullets(
        doc,
        [
            "Tableau-дашборд «Demand Planning»: drill-down по категорії / магазину / SKU, "
            "порівняння факт vs прогноз, відстеження bias та accuracy.",
            "Експорт CSV з прогнозом — для оффлайн-роботи планувальників.",
            "Email-нотифікація щопонеділка зі зведенням за тиждень.",
        ],
    )

    # ------------------------------------------------------------------ 19
    add_heading(doc, "19. Ризики")
    add_bullets(
        doc,
        [
            "Дрейф попиту (concept drift): сезонні зміни поведінки покупців — мітигація: "
            "квартальне перенавчання + автоматичний моніторинг WAPE.",
            "Зміна схеми DWH або відмова джерела даних — мітигація: data contracts + alerts у Airflow.",
            "Нові механіки промо (наприклад, flash-sale без попередження) — мітигація: "
            "ручний overrides від Category Manager у UI планувальника.",
            "Помилка в кат-фічах при появі нового SKU без історії — мітигація: fallback на "
            "категорійне середнє + флаг новизни SKU.",
            "Інфраструктурні: відмова Airflow worker, переповнення BigQuery quota — "
            "стандартні DevOps alerts + retry policy.",
        ],
    )

    # ------------------------------------------------------------------ 20
    add_heading(doc, "20. Примітки")
    add_bullets(
        doc,
        [
            "Поточна реалізація використовує синтетичний датасет (15 SKU × 2 магазини × 366 днів) "
            "як proof-of-concept; для продакшну треба підключити повні DWH-таблиці.",
            "Альтернативні моделі для майбутніх ітерацій: Prophet (per SKU), SARIMA, TFT/N-BEATS.",
            "Розглянути перехід з рекурсивного multi-step на direct multi-step "
            "(окремі моделі на +1, +2, +3, +4 тижні) — уникнення накопичення похибки.",
            "Документація розробника знаходиться у docs/REPORT.md та notebook/sku_sales_forecast.ipynb.",
        ],
    )

    # Збереження
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved to {OUT_PATH}")


if __name__ == "__main__":
    build()
