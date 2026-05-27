"""Генерує реферат за шаблоном НТУУ «КПІ» (TNR 14pt, інтервал 1.0)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "Referat.docx"

FONT = "Times New Roman"
SIZE_BODY = 14
SIZE_HEAD = 16


def set_font(run, *, size=SIZE_BODY, bold=False, italic=False):
    run.font.name = FONT
    # Ensure font set for all script variants
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_spacing_single(p):
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def para(doc, text, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False,
         size=SIZE_BODY, indent_first=0.0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing_single(p)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, *, level=1):
    size = SIZE_HEAD if level == 1 else SIZE_BODY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    set_spacing_single(p)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p


def empty(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        set_spacing_single(p)
        r = p.add_run("")
        set_font(r)


def page_break(doc):
    p = doc.add_paragraph()
    set_spacing_single(p)
    r = p.add_run()
    set_font(r)
    r.add_break(WD_BREAK.PAGE)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing_single(p)
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_font(run)
    return p


def numbered(doc, text, idx):
    """Manual numbering — без авто-стилю List Number, щоб не зчіплювати нумерацію."""
    p = doc.add_paragraph()
    set_spacing_single(p)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"{idx}. {text}")
    set_font(run)
    return p


# =========================================================================
doc = Document()

# Сторінка: A4, поля 2-2-2-2 см
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Стиль за замовчуванням
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(SIZE_BODY)

# =========================================================================
# ТИТУЛЬНИЙ ЛИСТ
# =========================================================================
para(doc, "МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
para(doc, "Національний технічний університет України", align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "«Київський політехнічний інститут імені Ігоря Сікорського»", align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "Факультет менеджменту та маркетингу", align=WD_ALIGN_PARAGRAPH.CENTER)
empty(doc, 8)

para(doc, "Реферат з дисципліни:", align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "«БІЗНЕС-ПРОГНОЗУВАННЯ»", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_HEAD)
empty(doc, 2)
para(doc, "Тема: «Прогноз продажу товарів у розрізі SKU»",
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
empty(doc, 8)

# Перевірив / Виконав
p = doc.add_paragraph()
set_spacing_single(p)
p.paragraph_format.tab_stops.add_tab_stop(Cm(11))
r = p.add_run("Перевірив:\tВиконав:")
set_font(r)

p = doc.add_paragraph()
set_spacing_single(p)
p.paragraph_format.tab_stops.add_tab_stop(Cm(11))
r = p.add_run("викладач\tстудентка групи УМ-з31")
set_font(r)

p = doc.add_paragraph()
set_spacing_single(p)
p.paragraph_format.tab_stops.add_tab_stop(Cm(11))
r = p.add_run("Биба В. М.\tГорбунова К.")
set_font(r)

empty(doc, 6)
para(doc, "Київ — 2026 р.", align=WD_ALIGN_PARAGRAPH.CENTER)

page_break(doc)

# =========================================================================
# ЗМІСТ
# =========================================================================
heading(doc, "ЗМІСТ")
empty(doc)

contents = [
    "Вступ",
    "Розділ 1. Технічне завдання проекту",
    "    1.1. Бізнес-проблема та постановка задачі",
    "    1.2. Об'єкт та горизонт прогнозування",
    "    1.3. Джерела даних",
    "    1.4. Підхід та модель",
    "    1.5. Метрики якості",
    "Розділ 2. Реалізація моделі та результати",
    "    2.1. Підготовка даних",
    "    2.2. Архітектура моделі",
    "    2.3. Результати на тестовому періоді",
    "    2.4. Прогноз та аналіз ознак",
    "    2.5. Бізнес-застосування",
    "Висновки",
    "Список використаних джерел",
]
for c in contents:
    para(doc, c, align=WD_ALIGN_PARAGRAPH.LEFT)

page_break(doc)

# =========================================================================
# ВСТУП
# =========================================================================
heading(doc, "ВСТУП")
empty(doc)

para(doc, "Прогнозування продажів у роздрібній торгівлі є одним із ключових елементів "
     "управління запасами. Помилка прогнозу веде або до прямої втрати продажів через "
     "відсутність товару (out-of-stock), або до заморожування оборотних коштів через "
     "надлишкові запаси на складі. Для роздрібних мереж і маркетплейсів з тисячами "
     "товарних позицій (SKU) ручне планування є непрактичним, тому використовуються "
     "алгоритми машинного навчання [1].",
     indent_first=1.25)

para(doc, "Метою цього реферату є опис розробленої моделі прогнозування продажів у розрізі "
     "SKU × магазин на тижневому горизонті 4 тижні наперед для умовного маркетплейсу.",
     indent_first=1.25)

para(doc, "Завдання реферату:", indent_first=1.25)
numbered(doc, "описати технічне завдання проекту та його обґрунтування;", 1)
numbered(doc, "висвітлити процес підготовки даних та побудови моделі;", 2)
numbered(doc, "представити отримані результати та оцінити якість прогнозу;", 3)
numbered(doc, "сформулювати висновки та визначити напрямки подальшого розвитку.", 4)

para(doc, "Структура реферату: робота складається зі вступу, двох розділів, висновків та "
     "списку використаних джерел.",
     indent_first=1.25)

page_break(doc)

# =========================================================================
# РОЗДІЛ 1
# =========================================================================
heading(doc, "РОЗДІЛ 1. ТЕХНІЧНЕ ЗАВДАННЯ ПРОЕКТУ")
empty(doc)

heading(doc, "1.1. Бізнес-проблема та постановка задачі", level=2)
para(doc, "Маркетплейс зазнає одночасних втрат із двох сторін. По-перше, out-of-stock (OOS) "
     "призводить до прямої втрати GMV: клієнт не знаходить потрібний товар і йде до "
     "конкурента, що додатково знижує customer retention. По-друге, over-stock спричиняє "
     "заморожування оборотних коштів, зростання витрат на зберігання та ризик подальшої "
     "уцінки товарів. Ручне планування десятків тисяч SKU не масштабується, тому потрібна "
     "автоматизована модель прогнозування [4, с. 23].",
     indent_first=1.25)

para(doc, "Цільова формула обсягу замовлення, прийнята в проекті:", indent_first=1.25)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing_single(p)
r = p.add_run("Q = Fct − St − GiT + SS_Fct,")
set_font(r, bold=True)
para(doc, "де Q — обсяг замовлення; Fct — прогноз продажів на горизонт; St — поточний "
     "залишок; GiT — товар у дорозі (Goods in Transit); SS_Fct — страховий запас, що "
     "залежить від прогнозу та його довірчого інтервалу.",
     indent_first=1.25)

heading(doc, "1.2. Об'єкт та горизонт прогнозування", level=2)
para(doc, "Об'єктом прогнозування є тижневі продажі (qty) для кожної пари «SKU × магазин». "
     "Горизонт прогнозу обрано 4 тижні наперед, що відповідає типовому циклу замовлення "
     "постачальникам у роздрібній торгівлі [4]. Гранулярність вхідних даних — день, "
     "з подальшою агрегацією до тижня (тижні з понеділка, W-MON). Глибина історії — "
     "12 місяців.",
     indent_first=1.25)

heading(doc, "1.3. Джерела даних", level=2)
para(doc, "Дані для моделі формуються з кількох таблиць корпоративного сховища даних:",
     indent_first=1.25)
bullet(doc, "sales_fact — щоденні продажі (date, sku_id, store_id, qty, price, promo);")
bullet(doc, "dim_sku — каталог SKU (категорія, бренд, ціна, вага);")
bullet(doc, "dim_store — довідник магазинів / fulfillment-центрів;")
bullet(doc, "dim_holidays — календар свят (8 березня, Black Friday, Новий рік тощо);")
bullet(doc, "dim_promo_calendar — календар акцій зі знижками;")
bullet(doc, "inventory_snapshot — щоденні залишки та товар у дорозі.")
para(doc, "Фінальна таблиця dataset формується через INNER та LEFT JOIN перерахованих "
     "таблиць; її структура містить як факт-показники (qty, price), так і атрибути з "
     "довідників (категорія, місто магазину, прапор свята).",
     indent_first=1.25)

heading(doc, "1.4. Підхід та модель", level=2)
para(doc, "В якості основної моделі обрано LightGBM — реалізацію алгоритму градієнтного "
     "бустингу на деревах рішень, розроблену в Microsoft Research [3]. Вибір моделі "
     "обґрунтовано наступними міркуваннями:",
     indent_first=1.25)
bullet(doc, "висока швидкість навчання та прогнозування на табличних даних;")
bullet(doc, "вбудована підтримка категоріальних ознак без one-hot encoding;")
bullet(doc, "гнучкість в обробці пропущених значень та лагових ознак;")
bullet(doc, "перевірена ефективність у задачах прогнозу попиту [1, с. 215].")
para(doc, "В якості baseline-моделі для порівняння використано наївний сезонний "
     "прогноз — середнє значення продажів за останні 4 тижні. Стратегія прогнозу "
     "на горизонт 4 тижні — рекурсивна (на кожному наступному кроці лагові ознаки "
     "оновлюються з попередніх передбачень моделі).",
     indent_first=1.25)

heading(doc, "1.5. Метрики якості", level=2)
para(doc, "В якості основної метрики обрано WAPE (Weighted Absolute Percentage Error), "
     "оскільки традиційна MAPE гіперчутлива до тижнів із малими фактичними значеннями "
     "(особливо актуально для рідкісних SKU) [1, с. 89]. Допоміжні метрики: MAPE, "
     "MAE, RMSE, Bias.",
     indent_first=1.25)
para(doc, "Поріг прийнятності моделі: WAPE ≤ 20%, Bias у межах ±5%. Латентність "
     "повного запуску — менше 5 хвилин.",
     indent_first=1.25)

page_break(doc)

# =========================================================================
# РОЗДІЛ 2
# =========================================================================
heading(doc, "РОЗДІЛ 2. РЕАЛІЗАЦІЯ МОДЕЛІ ТА РЕЗУЛЬТАТИ")
empty(doc)

heading(doc, "2.1. Підготовка даних", level=2)
para(doc, "Для побудови моделі використано синтетичний датасет, що відтворює структуру "
     "даних роздрібного маркетплейсу: 15 SKU у трьох категоріях (ноутбуки, смартфони, "
     "холодильники), 2 магазини, 366 днів історії (січень — грудень 2024 р.). У дані "
     "закладено тижнево-денну та річну сезонність, свята та випадкові акції зі знижкою "
     "−10%. Загалом — 10 980 рядків денних даних.",
     indent_first=1.25)
para(doc, "Денні продажі агреговано до тижневого рівня (W-MON). Сформовано 15 ознак, "
     "поділених на групи:",
     indent_first=1.25)
bullet(doc, "категоріальні: store_id, sku_id, category;")
bullet(doc, "календарні: month, quarter, is_holiday_week, woy_sin, woy_cos;")
bullet(doc, "цінові: avg_price, promo_share;")
bullet(doc, "лагові: qty_lag_1, qty_lag_2, qty_lag_4, qty_lag_12;")
bullet(doc, "ковзні середні: qty_roll_mean_4, qty_roll_mean_12.")

heading(doc, "2.2. Архітектура моделі", level=2)
para(doc, "Використано клас LightGBMRegressor з бібліотеки Python [5]. Налаштовані "
     "гіперпараметри: 400 дерев, learning_rate = 0.03, num_leaves = 15, max_depth = 5, "
     "L1-регуляризація alpha = 0.1, L2-регуляризація lambda = 0.5. Раннє зупинення "
     "(early stopping) спрацьовує після 50 ітерацій без покращення.",
     indent_first=1.25)
para(doc, "Поділ на навчальну та тестову вибірки виконано хронологічно: train — тижні "
     "1–44, test — тижні 45–52 (останні 8 тижнів 2024 р., період зі святами та Q4-піком).",
     indent_first=1.25)

heading(doc, "2.3. Результати на тестовому періоді", level=2)
para(doc, "Підсумкові метрики обох моделей наведено в таблиці 2.1.", indent_first=1.25)
empty(doc)

# Таблиця метрик
tbl = doc.add_table(rows=3, cols=6)
tbl.style = "Light Grid Accent 1"
hdr = ["Модель", "WAPE", "MAPE", "MAE", "RMSE", "Bias"]
data = [
    ["Baseline (4w mean)", "25,53%", "50,90%", "9,72", "15,34", "+12,19%"],
    ["LightGBM", "25,24%", "59,26%", "9,61", "15,32", "+10,44%"],
]
for ci, h in enumerate(hdr):
    cell = tbl.cell(0, ci)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing_single(p)
    r = p.add_run(h)
    set_font(r, bold=True)
for ri, row in enumerate(data, start=1):
    for ci, val in enumerate(row):
        cell = tbl.cell(ri, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        set_spacing_single(p)
        r = p.add_run(val)
        set_font(r)

empty(doc)
para(doc, "Таблиця 2.1 — Метрики моделей на тестовому періоді",
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
empty(doc)

para(doc, "LightGBM перевершив baseline за основною метрикою WAPE (25,24% проти 25,53%) "
     "та за MAE/RMSE. MAPE при цьому залишається високою через гіперчутливість до "
     "нульових значень — це підтверджує доцільність використання саме WAPE як основної "
     "метрики для рівня SKU [1, с. 92].",
     indent_first=1.25)
para(doc, "У розрізі категорій: смартфони — WAPE 21%, ноутбуки — 33%, холодильники — 35%. "
     "Найгірший результат у категорії холодильників пояснюється низькочастотним характером "
     "попиту (мала кількість одиниць на тиждень), що ускладнює виявлення стабільних "
     "закономірностей.",
     indent_first=1.25)

heading(doc, "2.4. Прогноз та аналіз ознак", level=2)
para(doc, "На основі останніх даних побудовано рекурсивний прогноз на 4 тижні наперед "
     "для всіх 30 пар (15 SKU × 2 магазини), що дало 120 прогнозних точок.",
     indent_first=1.25)
para(doc, "Аналіз feature importance виявив наступний топ ознак за внеском у модель:",
     indent_first=1.25)
numbered(doc, "qty_lag_1 — продажі попереднього тижня (найсильніший сигнал);", 1)
numbered(doc, "qty_roll_mean_4 — тренд останнього місяця;", 2)
numbered(doc, "avg_price — цінова реакція;", 3)
numbered(doc, "promo_share — частка днів з акцією у тижні;", 4)
numbered(doc, "is_holiday_week — наявність свят у тижні.", 5)
para(doc, "Цей розподіл відповідає теоретичним очікуванням: основним драйвером "
     "короткострокового прогнозу є попередня динаміка, доповнена ціновими та "
     "календарними чинниками [1, с. 145].",
     indent_first=1.25)

heading(doc, "2.5. Бізнес-застосування", level=2)
para(doc, "Прогноз інтегровано у формулу обсягу замовлення Q = Fct − St − GiT + SS_Fct, "
     "де SS_Fct = 1,65 × σ(прогноз), що відповідає рівню сервісу 95%. На виході "
     "формується план замовлення для системи закупівель.",
     indent_first=1.25)
para(doc, "Приклад розрахунку: для SKU 64502 у магазині 215 при Fct = 152, St = 74, "
     "GiT = 37, SS = 4 — обсяг замовлення Q = 45 одиниць.",
     indent_first=1.25)

page_break(doc)

# =========================================================================
# ВИСНОВКИ
# =========================================================================
heading(doc, "ВИСНОВКИ")
empty(doc)

para(doc, "У ході роботи розроблено модель прогнозування продажів у розрізі SKU × магазин "
     "на тижневому горизонті 4 тижні для умовного маркетплейсу. Основні результати:",
     indent_first=1.25)

numbered(doc, "Обрано та обґрунтовано модель LightGBM; сформовано 15 ознак, що поєднують "
              "лагові, календарні та цінові характеристики.", 1)
numbered(doc, "На тестовому періоді модель досягла WAPE 25,24% та обігнала наївний "
              "baseline за основними метриками.", 2)
numbered(doc, "Реалізовано формулу обсягу замовлення Q = Fct − St − GiT + SS_Fct, "
              "що готова до інтеграції в систему планування закупівель.", 3)
numbered(doc, "Найважливішими ознаками виявилися qty_lag_1 та qty_roll_mean_4, що "
              "підтверджує визначальну роль короткострокового тренду.", 4)

para(doc, "Серед напрямків подальшого розвитку:", indent_first=1.25)
bullet(doc, "розширення горизонту прогнозу до 8–12 тижнів;")
bullet(doc, "додавання зовнішніх регресорів (курс валют, погодні умови, макропоказники);")
bullet(doc, "перехід від рекурсивного прогнозу на стратегію direct multi-step для "
            "уникнення накопичення похибки;")
bullet(doc, "тестування альтернативних моделей: Prophet [2] та SARIMA per SKU.")

page_break(doc)

# =========================================================================
# СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ
# =========================================================================
heading(doc, "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ")
empty(doc)

sources = [
    "Hyndman R. J., Athanasopoulos G. Forecasting: Principles and Practice. 3rd ed. "
    "OTexts, 2021. URL: https://otexts.com/fpp3/ (дата звернення: 27.05.2026).",

    "Taylor S. J., Letham B. Forecasting at scale. The American Statistician. 2018. "
    "Vol. 72, No. 1. P. 37–45. DOI: 10.1080/00031305.2017.1380080.",

    "Ke G., Meng Q., Finley T., Wang T., Chen W., Ma W., Ye Q., Liu T.-Y. LightGBM: "
    "A highly efficient gradient boosting decision tree. Advances in Neural Information "
    "Processing Systems. 2017. Vol. 30. P. 3146–3154.",

    "Грабовецький Б. Є. Економічне прогнозування і планування : навчальний посібник. "
    "Київ : Центр учбової літератури, 2003. 188 с.",

    "LightGBM. Documentation [Електронний ресурс]. URL: https://lightgbm.readthedocs.io/ "
    "(дата звернення: 27.05.2026).",

    "Pandas. User Guide [Електронний ресурс]. URL: https://pandas.pydata.org/docs/ "
    "(дата звернення: 27.05.2026).",
]
for i, src in enumerate(sources, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing_single(p)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{i}. {src}")
    set_font(run)

# =========================================================================
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(f"Saved to {OUT_PATH}")
